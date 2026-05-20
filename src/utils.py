"""
utils.py — Report generation utilities for Explainova.

Requirements (install if not present):
    pip install python-docx
"""

from io import BytesIO
from datetime import datetime
import pandas as pd


def _build_preprocessing_steps(report: dict) -> list:
    steps = []

    removed_dupes = report.get("removed_duplicates", 0)
    if removed_dupes > 0:
        steps.append(f"Removed {removed_dupes} duplicate rows.")

    dropped_empty = report.get("dropped_empty_columns", [])
    if dropped_empty:
        steps.append(f"Dropped {len(dropped_empty)} fully-empty columns: {', '.join(dropped_empty)}")

    dropped_missing = report.get("dropped_high_missing_columns", [])
    if dropped_missing:
        steps.append(f"Dropped {len(dropped_missing)} high-missing columns (>60% missing).")

    dropped_single = report.get("dropped_single_value_columns", [])
    if dropped_single:
        steps.append(f"Dropped {len(dropped_single)} single-value (constant) columns.")

    dropped_id = report.get("dropped_id_columns", [])
    if dropped_id:
        steps.append(f"Dropped ID-like columns: {', '.join(dropped_id)}")

    dropped_hc = report.get("dropped_high_cardinality_columns", [])
    if dropped_hc:
        steps.append(f"Dropped {len(dropped_hc)} high-cardinality categorical columns.")

    rows_removed = report.get("removed_rows_due_to_missing", 0)
    if rows_removed > 0:
        steps.append(f"Removed {rows_removed} rows that exceeded the missing-value threshold.")

    parsed_dt = report.get("parsed_datetime_columns", [])
    if parsed_dt:
        steps.append(
            f"Parsed {len(parsed_dt)} datetime column(s) into year/month/day/weekday features: "
            f"{', '.join(parsed_dt)}"
        )

    zero_counts = report.get("zero_as_missing_counts", {})
    if zero_counts:
        replaced_total = sum(zero_counts.values())
        steps.append(
            f"Converted {replaced_total} zero placeholder value(s) to missing across "
            f"{len(zero_counts)} selected column(s)."
        )

    filled_num = report.get("filled_missing_numerical", [])
    if filled_num:
        steps.append(f"Filled missing values in {len(filled_num)} numeric column(s) using the column median.")

    filled_cat = report.get("filled_missing_categorical", [])
    if filled_cat:
        steps.append(f"Filled missing values in {len(filled_cat)} categorical column(s) using the mode.")

    capped = report.get("capped_outlier_columns", [])
    if capped:
        steps.append(f"Capped extreme outliers (3x IQR) in {len(capped)} column(s).")

    ordinal_enc = report.get("ordinal_encoded_columns", [])
    if ordinal_enc:
        steps.append(f"Ordinal-encoded columns: {', '.join(ordinal_enc)}")

    ohe = report.get("one_hot_encoded_columns", [])
    if ohe:
        steps.append(f"One-hot encoded {len(ohe)} remaining categorical column(s).")

    if report.get("target_encoded"):
        classes = report.get("target_classes", [])
        steps.append(f"Target variable label-encoded. Classes: {', '.join(str(c) for c in classes)}")

    if report.get("feature_reduction_applied"):
        low_var = len(report.get("removed_low_variance_columns", []))
        high_corr = len(report.get("removed_high_correlation_columns", []))
        steps.append(
            f"Feature selection applied with Variance Threshold and Pairwise Correlation - "
            f"removed {low_var} low-variance and {high_corr} highly-correlated feature(s)."
        )

    return steps


def _fmt(val) -> str:
    if isinstance(val, float):
        return f"{val:.4f}"
    return str(val)


def _get_primary_model_metric(problem_type: str):
    if problem_type == "classification":
        return "Accuracy", False
    return "R2 Score", False


def _get_best_model_summary(results_df: pd.DataFrame, problem_type: str):
    if results_df is None or results_df.empty:
        return None

    metric, ascending = _get_primary_model_metric(problem_type)
    if metric not in results_df.columns:
        return None

    best_row = results_df.sort_values(by=metric, ascending=ascending).iloc[0]
    return {
        "model_name": str(best_row["Model"]),
        "metric_name": metric,
        "metric_value": float(best_row[metric])
    }


def _add_table_explanation_lines(problem_type: str):
    if problem_type == "classification":
        return [
            "How to read this table:",
            "- Higher Accuracy, Precision, Recall, and F1 Score usually indicate better overall performance.",
            "- If ROC AUC is available, a higher value means cleaner class separation.",
            "- The best model is not always just the one with the highest accuracy; the type of error that matters most to you also matters."
        ]
    return [
        "How to read this table:",
        "- Higher R2 Score means the model explains more of the target variation.",
        "- Lower MAE and RMSE mean smaller prediction errors.",
        "- RMSE penalizes larger mistakes more strongly than MAE."
    ]


def _add_dataframe_table(doc, df: pd.DataFrame, title: str = None):
    if df is None or df.empty:
        return

    if title:
        doc.add_paragraph(title)

    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"

    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            table.rows[i + 1].cells[j].text = _fmt(row[col])

    doc.add_paragraph()


def _add_figure_to_doc(doc, fig, width_inches=5.8):
    from docx.shared import Inches

    if fig is None:
        return

    img_buf = BytesIO()
    fig.savefig(img_buf, format="png", bbox_inches="tight", dpi=130)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(width_inches))


def generate_word_report(
        target_column: str,
        report: dict,
        results_df: pd.DataFrame,
        problem_type: str,
        shap_outputs: dict = None,
        shap_model_name: str = None,
        shap_bar_fig=None,
        shap_summary_fig=None,
        corr_table: pd.DataFrame = None,
        corr_heatmap_fig=None,
        shap_effect_fig=None,
        effect_note: str = None,
        model_leaderboard_fig=None,
        model_recommendation_text: str = None,
        corr_profile_fig=None,
        corr_profile_note: str = None,
        feature_behavior_df: pd.DataFrame = None,
        feature_behavior_fig=None,
        kfold_df: pd.DataFrame = None,
        kfold_fig=None,
        kfold_note: str = None,
        pdp_ice_fig=None,
        pdp_ice_note: str = None,
        feature_detail_reports: list = None,
) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    INDIGO = RGBColor(0x4F, 0x46, 0xE5)
    SLATE = RGBColor(0x33, 0x41, 0x55)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = SLATE

    title_para = doc.add_heading("", level=0)
    title_run = title_para.add_run("Explainova - Analysis Report")
    title_run.font.color.rgb = INDIGO
    title_run.font.size = Pt(22)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
        f"Target: {target_column}   |   "
        f"Problem type: {problem_type.capitalize()}"
    ).font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ── Dataset Information ────────────────────────────────────────────────
    doc.add_heading("Dataset Overview", level=1)
    initial = report.get("initial_shape", ("N/A", "N/A"))
    final = report.get("final_shape", ("N/A", "N/A"))

    dataset_note = doc.add_paragraph()
    dataset_note.add_run(
        "This section compares the dataset before and after preprocessing. "
        "It helps you quickly see how much the cleaning and transformation steps changed the working data."
    )

    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Light Grid Accent 1"
    data_rows = [
        ("Initial shape", f"{initial[0]} rows x {initial[1]} columns"),
        ("Final shape (after preprocessing)", f"{final[0]} rows x {final[1]} columns"),
        ("Duplicate rows removed", str(report.get("removed_duplicates", 0))),
    ]
    for i, (label, value) in enumerate(data_rows):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── Preprocessing Summary ──────────────────────────────────────────────
    doc.add_heading("Preprocessing Summary", level=1)
    intro_p = doc.add_paragraph()
    intro_p.add_run(
        "The list below summarizes the main data-preparation steps applied before model training. "
        "These steps directly affect both model quality and how easy the results are to interpret."
    )

    for step in _build_preprocessing_steps(report):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(step)

    doc.add_paragraph()

    # ── Feature Relationship Overview ──────────────────────────────────────
    if corr_table is not None and not corr_table.empty:
        doc.add_heading("Feature Relationship Summary", level=1)

        corr_intro = doc.add_paragraph()
        corr_intro.add_run(
            "This section summarizes the strongest linear relationships between processed features and the target. "
            "Its purpose is to quickly highlight the variables that move most clearly with the target."
        )

        top_corr_df = corr_table[["Feature", "Correlation with Target"]].copy()
        _add_dataframe_table(doc, top_corr_df)

        if corr_profile_note:
            corr_story = doc.add_paragraph()
            corr_story.add_run(corr_profile_note)

        if corr_profile_fig is not None:
            _add_figure_to_doc(doc, corr_profile_fig, width_inches=5.8)

        if corr_heatmap_fig is not None:
            corr_note = doc.add_paragraph()
            corr_note.add_run(
                "The heatmap shows how the highlighted features move with one another and with the target. "
                "Values closer to 1 or -1 indicate stronger linear relationships, while values near 0 indicate weaker ones."
            )
            _add_figure_to_doc(doc, corr_heatmap_fig, width_inches=5.8)

        doc.add_paragraph()

    # ── Model Results ──────────────────────────────────────────────────────
    doc.add_heading("Model Results", level=1)

    summary = _get_best_model_summary(results_df, problem_type)
    if summary is not None:
        best_p = doc.add_paragraph()
        best_p.add_run("Leading model in the overall table: ").bold = True
        best_p.add_run(
            f"{summary['model_name']} based on {summary['metric_name']} = {summary['metric_value']:.4f}"
        )

    if model_recommendation_text:
        recommendation_p = doc.add_paragraph()
        recommendation_p.add_run(model_recommendation_text)

    if model_leaderboard_fig is not None:
        leaderboard_note = doc.add_paragraph()
        leaderboard_note.add_run(
            "This chart ranks the models by the main success metric at a glance. "
            "The model at the top is the strongest candidate for a first recommendation."
        )
        _add_figure_to_doc(doc, model_leaderboard_fig, width_inches=5.9)

    for line in _add_table_explanation_lines(problem_type):
        p = doc.add_paragraph(style="List Bullet" if line.startswith("-") else None)
        p.add_run(line[2:] if line.startswith("-") else line)

    if results_df is not None and not results_df.empty:
        _add_dataframe_table(doc, results_df)

    if kfold_df is not None and not kfold_df.empty:
        doc.add_heading("Stability Check with K-Fold", level=2)
        kfold_intro = doc.add_paragraph()
        kfold_intro.add_run(
            "K-fold validation repeats evaluation across several train/test splits. "
            "It helps determine whether the model result is stable across different data partitions."
        )
        if kfold_note:
            doc.add_paragraph(kfold_note)
        if kfold_fig is not None:
            _add_figure_to_doc(doc, kfold_fig, width_inches=5.9)
        compact_kfold_cols = [
            col for col in kfold_df.columns
            if col in [
                "Model", "Folds",
                "ROC AUC Mean", "ROC AUC Std",
                "F1 Score Mean", "F1 Score Std",
                "Accuracy Mean", "Accuracy Std",
                "R2 Score Mean", "R2 Score Std",
                "MAE Mean", "MAE Std",
                "RMSE Mean", "RMSE Std",
            ]
        ]
        _add_dataframe_table(doc, kfold_df[compact_kfold_cols] if compact_kfold_cols else kfold_df)

    # ── SHAP Explainability ────────────────────────────────────────────────
    if shap_outputs is not None:
        doc.add_heading(f"SHAP Explanations - {shap_model_name}", level=1)

        intro = doc.add_paragraph()
        intro.add_run(
            "SHAP shows which features influenced a model result and by how much. "
            "Positive values push the result upward, while negative values pull it downward. "
            "Larger absolute values mean stronger influence."
        )

        importance_df = shap_outputs.get("feature_importance_df")
        if importance_df is not None and not importance_df.empty:
            doc.add_heading("Top SHAP Features", level=2)

            shap_note = doc.add_paragraph()
            shap_note.add_run(
                "This table ranks features by average impact strength. "
                "Items near the top influence the result more strongly and more consistently across the analyzed samples."
            )

            top_df = importance_df.head(8).copy()
            _add_dataframe_table(doc, top_df)

        if shap_bar_fig is not None:
            doc.add_heading("Impact Strength Chart", level=2)
            chart_note = doc.add_paragraph()
            chart_note.add_run(
                "Longer bars mean the feature has a stronger overall effect on the model result."
            )
            _add_figure_to_doc(doc, shap_bar_fig, width_inches=5.8)

        if shap_summary_fig is not None:
            doc.add_heading("Overall SHAP Distribution", level=2)
            summary_note = doc.add_paragraph()
            summary_note.add_run(
                "Each dot represents one sample. A dot's position to the right or left shows whether that feature pushed the result upward or downward."
            )
            _add_figure_to_doc(doc, shap_summary_fig, width_inches=6.1)

        if feature_detail_reports:
            doc.add_heading("Selected Feature Behavior", level=2)
            feature_intro = doc.add_paragraph()
            feature_intro.add_run(
                "This section includes the feature-level behavior charts selected in the application. "
                "The SHAP effect chart shows observed contribution direction, while PDP / ICE shows controlled model response when the feature value changes."
            )

            for item in feature_detail_reports:
                feature_name = item.get("feature_name", "Selected feature")
                doc.add_heading(str(feature_name), level=3)

                effect_item_note = item.get("effect_note")
                if effect_item_note:
                    doc.add_paragraph(effect_item_note)
                if item.get("effect_fig") is not None:
                    _add_figure_to_doc(doc, item.get("effect_fig"), width_inches=5.3)

                pdp_item_note = item.get("pdp_ice_note")
                if pdp_item_note:
                    doc.add_paragraph(pdp_item_note)
                if item.get("pdp_ice_fig") is not None:
                    _add_figure_to_doc(doc, item.get("pdp_ice_fig"), width_inches=5.3)
        else:
            if shap_effect_fig is not None:
                doc.add_heading("How One Feature Changes the Result", level=2)

                intro_effect = doc.add_paragraph()
                intro_effect.add_run(
                    "This chart shows how the model usually reacts when one feature changes. "
                    "It helps reveal whether larger or smaller values tend to increase or decrease the result."
                )

                if effect_note:
                    p = doc.add_paragraph()
                    p.add_run(effect_note)

                _add_figure_to_doc(doc, shap_effect_fig, width_inches=5.5)

            if pdp_ice_fig is not None:
                doc.add_heading("PDP / ICE Feature Movement", level=2)
                pdp_intro = doc.add_paragraph()
                pdp_intro.add_run(
                    "PDP and ICE actively change one feature and show how the model response moves. "
                    "This complements SHAP by showing model behavior under controlled what-if changes."
                )
                if pdp_ice_note:
                    doc.add_paragraph(pdp_ice_note)
                _add_figure_to_doc(doc, pdp_ice_fig, width_inches=5.8)

        if feature_behavior_fig is not None:
            doc.add_heading("Overall Behavior of Key Features", level=2)
            feature_behavior_note = doc.add_paragraph()
            feature_behavior_note.add_run(
                "This chart summarizes the average impact strength and typical direction of the most important features. "
                "It helps you quickly see which features behave more strongly and more consistently."
            )
            _add_figure_to_doc(doc, feature_behavior_fig, width_inches=5.9)

        if feature_behavior_df is not None and not feature_behavior_df.empty:
            _add_dataframe_table(doc, feature_behavior_df.head(6), title="Behavior summary table")

    # ── Footer ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Generated by Explainova | Powered by SHAP and scikit-learn")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
